#!/usr/bin/env python
from __future__ import annotations

import argparse
import os
import json
import random
import re
import shutil
from collections import Counter
from datetime import date, timedelta
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE_ROOT = Path(os.environ.get("PROVDATA_SOURCE_ROOT", ROOT / "data/raw/user_uploads"))


LAW_SLUGS = {
    "中华人民共和国农村土地承包法": "Rural_Land_Contract_Law",
    "中华人民共和国农村土地承包经营纠纷调解仲裁法": "Rural_Land_Contract_Arbitration_Law",
    "中华人民共和国土地管理法": "Land_Administration_Law",
    "中华人民共和国土地管理法实施条例": "Land_Administration_Regulation",
    "中华人民共和国民法典": "Civil_Code",
    "中华人民共和国环境保护法": "Environmental_Protection_Law",
    "中华人民共和国草原法": "Grassland_Law",
    "最高人民法院关于审理涉及农村土地承包经营纠纷调解仲裁案件适用法律若干问题的解释": "Rural_Land_Contract_Dispute_Judicial_Interpretation",
    "草原防火条例": "Grassland_Fire_Prevention_Regulation",
}


EFFECTIVE_OVERRIDES = {
    ("中华人民共和国民法典", "20200528"): "2021-01-01",
    ("中华人民共和国环境保护法", "20140424"): "2015-01-01",
    ("中华人民共和国土地管理法实施条例", "20210702"): "2021-09-01",
    ("中华人民共和国农村土地承包经营纠纷调解仲裁法", "20090627"): "2010-01-01",
}


ARTICLE_RE = re.compile(r"(第[一二三四五六七八九十百千万零〇两]+条)\s*")


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    chunks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" ".join(cells))
    return "\n".join(chunks)


def parse_law_filename(path: Path) -> tuple[str, str]:
    stem = path.stem
    match = re.match(r"(.+)_([0-9]{8})$", stem)
    if not match:
        return stem, "19000101"
    return match.group(1), match.group(2)


def ymd(raw: str) -> str:
    return f"{raw[:4]}-{raw[4:6]}-{raw[6:8]}"


def slug_for_law(law_name: str) -> str:
    return LAW_SLUGS.get(law_name, re.sub(r"\W+", "_", law_name).strip("_"))


def split_articles(text: str) -> list[tuple[str, str]]:
    matches = list(ARTICLE_RE.finditer(text))
    if not matches:
        compact = re.sub(r"\s+", " ", text).strip()
        return [("全文", compact)] if compact else []
    articles = []
    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        article_no = match.group(1)
        article_text = re.sub(r"\s+", " ", text[start:end]).strip()
        if len(article_text) >= 8:
            articles.append((article_no, article_text))
    return articles


def build_law_records(law_doc_dir: Path) -> list[dict]:
    source_files = sorted(p for p in law_doc_dir.glob("*.docx") if not p.name.startswith("~$"))
    version_meta = []
    for path in source_files:
        law_name, version_raw = parse_law_filename(path)
        version_meta.append(
            {
                "path": path,
                "law_name": law_name,
                "version_raw": version_raw,
                "version_date": ymd(version_raw),
                "valid_from": EFFECTIVE_OVERRIDES.get((law_name, version_raw), ymd(version_raw)),
            }
        )

    by_law: dict[str, list[dict]] = {}
    for item in version_meta:
        by_law.setdefault(item["law_name"], []).append(item)
    for versions in by_law.values():
        versions.sort(key=lambda x: x["valid_from"])
        for idx, item in enumerate(versions):
            if idx + 1 < len(versions):
                item["valid_to"] = (date.fromisoformat(versions[idx + 1]["valid_from"]) - timedelta(days=1)).isoformat()
                item["status"] = "expired"
            else:
                item["valid_to"] = None
                item["status"] = "effective"

    records = []
    for item in version_meta:
        law_slug = slug_for_law(item["law_name"])
        text = read_docx_text(item["path"])
        articles = split_articles(text)
        for article_idx, (article_no, article_text) in enumerate(articles, start=1):
            article_slug = f"A{article_idx:03d}" if article_no == "全文" else f"A{article_idx:03d}"
            law_id = f"{law_slug}_{item['version_raw']}_{article_slug}"
            records.append(
                {
                    "law_id": law_id,
                    "law_key": f"{law_slug}|{article_no}",
                    "law_name": item["law_name"],
                    "article_no": article_no,
                    "article_text": article_text,
                    "article_summary": article_text[:120],
                    "risk_tags": [],
                    "article_nature": "real_law_article",
                    "version": item["version_raw"],
                    "version_date": item["version_date"],
                    "t_start": item["valid_from"],
                    "t_end": item["valid_to"],
                    "valid_from": item["valid_from"],
                    "valid_to": item["valid_to"],
                    "status": item["status"],
                    "source_file": item["path"].name,
                    "source_type": "real_docx",
                }
            )
    return records


def write_jsonl(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def load_jsonl(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8") as f:
        return [json.loads(line) for line in f if line.strip()]


def normalize_annotations(source: Path) -> list[dict]:
    rows = load_jsonl(source)
    normalized = []
    for row in rows:
        label = int(row.get("risk_label", 0))
        anchor = row.get("time_anchor") or row.get("contract_effective_date") or row.get("contract_sign_date")
        normalized.append(
            {
                "sample_id": row.get("sample_id"),
                "contract_id": row.get("contract_id"),
                "clause_id": row.get("clause_id"),
                "clause_no": row.get("clause_no"),
                "dataset": "GrassRiskReal",
                "task": "contract_risk_review",
                "contract_type": row.get("contract_type"),
                "source_file": row.get("source_file"),
                "contract_sign_date": row.get("contract_sign_date"),
                "contract_effective_date": row.get("contract_effective_date"),
                "anchor_date": anchor,
                "time_anchor": anchor,
                "clause_text": row.get("clause_text", ""),
                "risk_type": row.get("risk_type", "unknown"),
                "label": label,
                "risk_label": label,
                "label_name": "风险条款" if label else "非风险条款",
                "temporal_label": int(row.get("temporal_label", 0)),
                "temporal_error_type": row.get("temporal_error_type"),
                "gold_evidence_ids": row.get("gold_basis_ids", []),
                "gold_basis_ids": row.get("gold_basis_ids", []),
                "gold_risk_judgment": row.get("gold_risk_judgment"),
                "gold_legal_analysis": row.get("gold_legal_analysis"),
                "gold_temporal_explanation": row.get("gold_temporal_explanation"),
                "gold_consistency_label": row.get("gold_consistency_label"),
                "review_steps": {
                    "evidence_summary": row.get("gold_legal_analysis", ""),
                    "clause_evidence_alignment": row.get("gold_legal_analysis", ""),
                    "temporal_consequence": row.get("gold_temporal_explanation", ""),
                },
                "source_type": "real_annotation",
            }
        )
    return normalized


def grouped_split(rows: list[dict], seed: int = 42) -> dict[str, list[dict]]:
    rng = random.Random(seed)
    contracts = sorted({row["contract_id"] for row in rows})
    rng.shuffle(contracts)
    n = len(contracts)
    train_contracts = set(contracts[: max(1, round(n * 0.6))])
    val_contracts = set(contracts[max(1, round(n * 0.6)) : max(1, round(n * 0.8))])
    splits = {"train": [], "val": [], "test": []}
    for row in rows:
        cid = row["contract_id"]
        if cid in train_contracts:
            splits["train"].append(row)
        elif cid in val_contracts:
            splits["val"].append(row)
        else:
            splits["test"].append(row)
    return splits


def copy_docx_files(source_dir: Path, target_dir: Path, text_dir: Path | None = None) -> list[str]:
    target_dir.mkdir(parents=True, exist_ok=True)
    if text_dir:
        text_dir.mkdir(parents=True, exist_ok=True)
    copied = []
    for path in sorted(source_dir.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        target = target_dir / path.name
        shutil.copy2(path, target)
        copied.append(path.name)
        if text_dir:
            text = read_docx_text(path)
            (text_dir / f"{path.stem}.txt").write_text(text, encoding="utf-8")
    return copied


def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest real grassland contracts, laws, and annotations.")
    parser.add_argument("--source-root", type=Path, default=DEFAULT_SOURCE_ROOT)
    parser.add_argument("--seed", type=int, default=42)
    args = parser.parse_args()

    contracts_src = args.source_root / "草原承包合同"
    laws_src = args.source_root / "草原承包合同有关法律"
    annotations_src = args.source_root / "grassrisk_annotations.jsonl"
    for path in [contracts_src, laws_src, annotations_src]:
        if not path.exists():
            raise FileNotFoundError(path)

    contract_files = copy_docx_files(contracts_src, ROOT / "data/raw/contracts", ROOT / "data/raw/contracts_text")
    law_files = copy_docx_files(laws_src, ROOT / "data/raw/laws/source_docs")

    law_records = build_law_records(laws_src)
    write_jsonl(ROOT / "data/raw/laws/legal_validity_kb.jsonl", law_records)
    write_jsonl(ROOT / "data/raw/legal_validity_kb.jsonl", law_records)

    raw_ann_target = ROOT / "data/raw/annotations/grassrisk_annotations.jsonl"
    raw_ann_target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(annotations_src, raw_ann_target)

    normalized = normalize_annotations(annotations_src)
    out_dir = ROOT / "data/processed/GrassRiskReal"
    write_jsonl(out_dir / "all.jsonl", normalized)
    splits = grouped_split(normalized, args.seed)
    for split_name, split_rows in splits.items():
        write_jsonl(out_dir / f"{split_name}.jsonl", split_rows)
    meta = {
        "dataset": "GrassRiskReal",
        "source_annotations": str(annotations_src),
        "total": len(normalized),
        "positive": sum(1 for row in normalized if row["label"] == 1),
        "negative": sum(1 for row in normalized if row["label"] == 0),
        "temporal_positive": sum(1 for row in normalized if row["temporal_label"] == 1),
        "contracts": len({row["contract_id"] for row in normalized}),
        "split_sizes": {name: len(rows) for name, rows in splits.items()},
        "risk_type_counts": dict(Counter(row["risk_type"] for row in normalized)),
        "source_type": "real_annotation",
    }
    (out_dir / "meta.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    report = {
        "contracts_copied": len(contract_files),
        "law_docs_copied": len(law_files),
        "law_articles": len(law_records),
        "annotations_normalized": len(normalized),
        "outputs": {
            "contracts": "data/raw/contracts",
            "contract_text": "data/raw/contracts_text",
            "law_docs": "data/raw/laws/source_docs",
            "law_jsonl": "data/raw/laws/legal_validity_kb.jsonl",
            "annotations": "data/processed/GrassRiskReal",
        },
    }
    report_path = ROOT / "outputs/real_data_ingest_summary.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
